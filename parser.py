import shutil
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from database import Document, Paragraph, Sentence, UPLOAD_DIR, Word
from splitter import normalize_text, split_paragraphs, split_sentences, split_words


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def validate_extension(file_name: str) -> str:
    extension = Path(file_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Танҳо Китобхои PDF ва DOCX дастгирӣ мешаванд.")
    return extension


def copy_file_to_storage(source_path: Path) -> tuple[str, Path]:
    extension = validate_extension(source_path.name)
    stored_name = f"{uuid4().hex}{extension}"
    destination = UPLOAD_DIR / stored_name
    shutil.copy2(source_path, destination)
    return extension, destination


def save_upload_file(upload_file) -> tuple[str, Path]:
    extension = validate_extension(upload_file.filename or "")
    stored_name = f"{uuid4().hex}{extension}"
    destination = UPLOAD_DIR / stored_name
    with destination.open("wb") as buffer:
        shutil.copyfileobj(upload_file.file, buffer)
    return extension, destination


def create_uploaded_document(
    db: Session,
    filename: str,
    file_type: str,
    stored_path: Path,
    *,
    title: str | None = None,
    author: str | None = None,
    publisher: str | None = None,
    publication_year: int | None = None,
    doc_type: str | None = None,
    bibliography: str | None = None,
) -> Document:
    document = Document(
        filename=filename,
        file_type=file_type,
        stored_path=str(stored_path),
        status="uploaded",
        title=title,
        author=author,
        publisher=publisher,
        publication_year=publication_year,
        doc_type=doc_type,
        bibliography=bibliography,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


def extract_pdf_text(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    page_texts = [page.extract_text() or "" for page in reader.pages]
    return normalize_text("\n\n".join(page_texts))


def extract_docx_text(file_path: Path) -> str:
    document = DocxDocument(str(file_path))
    paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return normalize_text("\n\n".join(paragraph_texts))


def extract_text(file_path: Path, file_type: str) -> str:
    if file_type == ".pdf":
        return extract_pdf_text(file_path)
    if file_type == ".docx":
        return extract_docx_text(file_path)
    raise ValueError("Формати файл дастгирӣ намешавад.")


def parse_document(db: Session, document: Document) -> dict[str, int]:
    try:
        structured_text = extract_text(Path(document.stored_path), document.file_type)
        if not structured_text:
            raise ValueError("Аз файл матн бароварда нашуд.")

        document.full_text = " ".join(structured_text.split())
        document.status = "parsed"
        document.error_message = None

        db.query(Word).filter(Word.document_id == document.id).delete()
        db.query(Sentence).filter(Sentence.document_id == document.id).delete()
        db.query(Paragraph).filter(Paragraph.document_id == document.id).delete()

        paragraphs = split_paragraphs(structured_text)
        sentence_counter = 0
        word_counter = 0

        for paragraph_index, paragraph_text in enumerate(paragraphs, start=1):
            paragraph = Paragraph(
                document_id=document.id,
                paragraph_index=paragraph_index,
                text=paragraph_text,
            )
            db.add(paragraph)
            db.flush()

            sentences = split_sentences(paragraph_text)
            for sentence_index, sentence_text in enumerate(sentences, start=1):
                sentence_counter += 1
                sentence = Sentence(
                    document_id=document.id,
                    paragraph_id=paragraph.id,
                    sentence_index=sentence_index,
                    text=sentence_text,
                )
                db.add(sentence)
                db.flush()

                for word_text in split_words(sentence_text):
                    word_counter += 1
                    db.add(
                        Word(
                            document_id=document.id,
                            sentence_id=sentence.id,
                            word_index=word_counter,
                            word=word_text,
                        )
                    )

        db.commit()
        return {
            "paragraphs": len(paragraphs),
            "sentences": sentence_counter,
            "words": word_counter,
        }
    except Exception as exc:
        db.rollback()
        document.status = "error"
        document.error_message = str(exc)
        db.add(document)
        db.commit()
        return {"paragraphs": 0, "sentences": 0, "words": 0}


def load_document_from_path(
    db: Session,
    file_path: str,
    *,
    title: str | None = None,
    author: str | None = None,
    publisher: str | None = None,
    publication_year: int | None = None,
    doc_type: str | None = None,
    bibliography: str | None = None,
) -> tuple[Document, dict[str, int]]:
    source_path = Path(file_path)
    if not source_path.exists() or not source_path.is_file():
        raise ValueError("Файл ёфт нашуд.")

    file_type, stored_path = copy_file_to_storage(source_path)
    document = create_uploaded_document(
        db,
        source_path.name,
        file_type,
        stored_path,
        title=title,
        author=author,
        publisher=publisher,
        publication_year=publication_year,
        doc_type=doc_type,
        bibliography=bibliography,
    )
    counts = parse_document(db, document)
    db.refresh(document)
    return document, counts
