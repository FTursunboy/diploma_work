from collections import Counter
import hashlib
import json
import shutil
import subprocess
import threading
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from database import (
    Document,
    DocumentChunk,
    DocumentNgram,
    Paragraph,
    ParagraphEmbeddingBlock,
    Sentence,
    SessionLocal,
    UPLOAD_DIR,
    Word,
)
from services.ai_log_service import AiLogService
from services.chunk_service import split_text_chunks
from services.embedding_service import EmbeddingService
from services.summary_service import SummaryService
from splitter import normalize_text, split_paragraphs, split_sentences, split_words


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}
NGRAM_MIN_N = 2
NGRAM_MAX_N = 5


def _ngram_hash(ngram: str) -> str:
    return hashlib.sha256(ngram.encode("utf-8")).hexdigest()


def _add_sentence_ngrams(tokens: list[str], counts: Counter[tuple[int, str]]) -> None:
    for n in range(NGRAM_MIN_N, NGRAM_MAX_N + 1):
        if len(tokens) < n:
            continue
        for index in range(0, len(tokens) - n + 1):
            ngram = " ".join(tokens[index : index + n])
            counts[(n, ngram)] += 1


class DocumentParserService:
    def __init__(self, db: Session | None):
        self._db = db

    def validate_extension(self, file_name: str) -> str:
        extension = Path(file_name).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError("Танҳо файлҳои PDF, DOC ва DOCX дастгирӣ мешаванд.")
        return extension

    def copy_file_to_storage(self, source_path: Path) -> tuple[str, Path]:
        extension = self.validate_extension(source_path.name)
        stored_name = f"{uuid4().hex}{extension}"
        destination = UPLOAD_DIR / stored_name
        shutil.copy2(source_path, destination)
        return extension, destination

    def save_upload_file(self, upload_file) -> tuple[str, Path]:
        extension = self.validate_extension(upload_file.filename or "")
        stored_name = f"{uuid4().hex}{extension}"
        destination = UPLOAD_DIR / stored_name
        with destination.open("wb") as buffer:
            shutil.copyfileobj(upload_file.file, buffer)
        return extension, destination

    def create_uploaded_document(
        self,
        *,
        filename: str,
        file_type: str,
        stored_path: Path,
        title: str | None = None,
        author: str | None = None,
        publisher: str | None = None,
        publication_year: int | None = None,
        doc_type: str | None = None,
        bibliography: str | None = None,
    ) -> Document:
        if self._db is None:
            raise RuntimeError("Database session is required for create_uploaded_document().")
        document = Document(
            filename=filename,
            file_type=file_type,
            stored_path=str(stored_path),
            status="uploaded",
            title=title,
            author=author,
            publisher=publisher,
            publication_year=publication_year,
            doc_type=doc_type,
            bibliography=bibliography,
        )
        self._db.add(document)
        self._db.commit()
        self._db.refresh(document)
        return document

    def extract_pdf_text(self, file_path: Path) -> str:
        reader = PdfReader(str(file_path))
        page_texts = [page.extract_text() or "" for page in reader.pages]
        return normalize_text("\n\n".join(page_texts))

    def extract_docx_text(self, file_path: Path) -> str:
        document = DocxDocument(str(file_path))
        paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        return normalize_text("\n\n".join(paragraph_texts))

    def _is_rtf_file(self, file_path: Path) -> bool:
        try:
            header = file_path.read_bytes()[:32]
        except OSError:
            return False
        return header.lstrip().startswith(b"{\\rtf")

    def extract_rtf_text(self, file_path: Path) -> str:
        try:
            result = subprocess.run(
                ["unrtf", "--text", str(file_path)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=True,
            )
        except FileNotFoundError as exc:
            raise RuntimeError("Утилита unrtf не установлена. Пересоберите Docker-образ.") from exc
        except subprocess.CalledProcessError as exc:
            error_text = (exc.stderr or exc.stdout or "").strip()
            raise RuntimeError(error_text or "Не удалось извлечь текст из RTF-файла.") from exc

        cleaned = "\n".join(
            line for line in (result.stdout or "").splitlines() if not line.startswith("### ")
        )
        return normalize_text(cleaned)

    def extract_doc_text(self, file_path: Path) -> str:
        if self._is_rtf_file(file_path):
            return self.extract_rtf_text(file_path)
        try:
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
                check=True,
            )
        except FileNotFoundError as exc:
            raise RuntimeError("Утилита antiword не установлена. Пересоберите Docker-образ.") from exc
        except subprocess.CalledProcessError as exc:
            error_text = (exc.stderr or exc.stdout or "").strip()
            if "Rich Text Format file" in error_text or "probably a Rich Text Format file" in error_text:
                return self.extract_rtf_text(file_path)
            raise RuntimeError(error_text or "Не удалось извлечь текст из DOC-файла.") from exc

        return normalize_text(result.stdout or "")

    def extract_text(self, file_path: Path, file_type: str) -> str:
        if file_type == ".pdf":
            return self.extract_pdf_text(file_path)
        if file_type == ".doc":
            return self.extract_doc_text(file_path)
        if file_type == ".docx":
            return self.extract_docx_text(file_path)
        raise ValueError("Формати файл дастгирӣ намешавад.")

    def parse_document(self, document: Document) -> dict[str, int]:
        if self._db is None:
            raise RuntimeError("Database session is required for parse_document().")
        try:
            structured_text = self.extract_text(Path(document.stored_path), document.file_type)
            if not structured_text:
                raise ValueError("Аз файл матн бароварда нашуд.")

            document.full_text = " ".join(structured_text.split())
            document.status = "parsed"
            document.ai_status = "pending"
            document.ai_summary = None
            document.error_message = None

            self._db.query(Word).filter(Word.document_id == document.id).delete()
            self._db.query(Sentence).filter(Sentence.document_id == document.id).delete()
            self._db.query(Paragraph).filter(Paragraph.document_id == document.id).delete()
            self._db.query(DocumentNgram).filter(DocumentNgram.document_id == document.id).delete()
            self._db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()

            paragraphs = split_paragraphs(structured_text)
            sentence_counter = 0
            word_counter = 0
            ngram_counts: Counter[tuple[int, str]] = Counter()

            for paragraph_index, paragraph_text in enumerate(paragraphs, start=1):
                paragraph = Paragraph(
                    document_id=document.id,
                    paragraph_index=paragraph_index,
                    text=paragraph_text,
                )
                self._db.add(paragraph)
                self._db.flush()

                sentences = split_sentences(paragraph_text)
                for sentence_index, sentence_text in enumerate(sentences, start=1):
                    sentence_counter += 1
                    sentence = Sentence(
                        document_id=document.id,
                        paragraph_id=paragraph.id,
                        sentence_index=sentence_index,
                        text=sentence_text,
                    )
                    self._db.add(sentence)
                    self._db.flush()

                    sentence_words = split_words(sentence_text)
                    ngram_tokens = [(word or "").strip().lower() for word in sentence_words if (word or "").strip()]
                    _add_sentence_ngrams(ngram_tokens, ngram_counts)

                    for word_text in sentence_words:
                        word_counter += 1
                        self._db.add(
                            Word(
                                document_id=document.id,
                                sentence_id=sentence.id,
                                word_index=word_counter,
                                word=word_text,
                            )
                        )

            for (n, ngram), count in ngram_counts.items():
                self._db.add(
                    DocumentNgram(
                        document_id=document.id,
                        n=n,
                        ngram=ngram,
                        ngram_hash=_ngram_hash(ngram),
                        count=int(count),
                    )
                )

            self._db.commit()

            return {
                "paragraphs": len(paragraphs),
                "sentences": sentence_counter,
                "words": word_counter,
                "chunks": 0,
            }
        except Exception as exc:
            self._db.rollback()
            document.status = "error"
            document.error_message = str(exc)
            self._db.add(document)
            self._db.commit()
            return {"paragraphs": 0, "sentences": 0, "words": 0, "chunks": 0}

    def run_ai_processing_for_document(self, *, document_id: int) -> None:
        if self._db is None:
            raise RuntimeError("Database session is required for run_ai_processing_for_document().")

        document = self._db.get(Document, document_id)
        if document is None or document.status != "parsed":
            return

        if not EmbeddingService.is_configured():
            message = "OPENAI_API_KEY is not configured."
            document.ai_status = "error"
            document.error_message = message
            self._db.add(document)
            self._db.commit()
            AiLogService(self._db).create(
                operation="ai_processing",
                status="error",
                document_id=document.id,
                error_message=message,
            )
            return

        document.ai_status = "processing"
        self._db.add(document)
        self._db.commit()

        try:
            structured_text = self.extract_text(Path(document.stored_path), document.file_type)
            if not structured_text:
                structured_text = document.full_text or ""
            self.index_document_paragraphs(document=document)
            self.index_document_chunks(document=document, structured_text=structured_text)
            SummaryService(self._db).generate_for_document(document=document, text=structured_text)
            document.ai_status = "ready"
            document.error_message = None
            self._db.add(document)
            self._db.commit()
        except Exception as exc:
            self._db.rollback()
            document = self._db.get(Document, document_id)
            if document is not None:
                document.ai_status = "error"
                document.error_message = str(exc)
                self._db.add(document)
                self._db.commit()

    def index_document_chunks(self, *, document: Document, structured_text: str) -> int:
        if self._db is None:
            raise RuntimeError("Database session is required for index_document_chunks().")

        chunks = split_text_chunks(structured_text)
        if not chunks:
            return 0

        self._db.query(DocumentChunk).filter(DocumentChunk.document_id == document.id).delete()

        embedding_service = None
        if EmbeddingService.is_configured():
            try:
                embedding_service = EmbeddingService()
            except Exception:
                embedding_service = None

        created = 0
        for chunk_index, chunk_text in enumerate(chunks, start=1):
            embedding_json = None
            embedding_model = None
            if embedding_service is not None:
                started_model = embedding_service.model
                try:
                    result = embedding_service.get_embedding_result(chunk_text)
                    embedding_json = json.dumps(result.embedding, separators=(",", ":"))
                    embedding_model = result.model
                    AiLogService(self._db).create(
                        operation="embedding_chunk",
                        status="success",
                        document_id=document.id,
                        model=result.model,
                        request_text=chunk_text,
                        response_text=f"chunk_index={chunk_index}; vector_dimensions={len(result.embedding)}",
                        prompt_tokens=result.prompt_tokens,
                        total_tokens=result.total_tokens,
                        duration_ms=result.duration_ms,
                    )
                except Exception as exc:
                    embedding_json = None
                    embedding_model = started_model
                    AiLogService(self._db).create(
                        operation="embedding_chunk",
                        status="error",
                        document_id=document.id,
                        model=started_model,
                        request_text=chunk_text,
                        error_message=str(exc),
                    )

            self._db.add(
                DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk_index,
                    chunk_text=chunk_text,
                    embedding=embedding_json,
                    embedding_model=embedding_model,
                )
            )
            created += 1

        self._db.commit()
        return created

    def index_document_paragraphs(self, *, document: Document) -> int:
        if self._db is None:
            raise RuntimeError("Database session is required for index_document_paragraphs().")

        paragraphs = self._db.scalars(
            select(Paragraph).where(Paragraph.document_id == document.id).order_by(Paragraph.paragraph_index)
        ).all()
        if not paragraphs:
            return 0

        self._db.query(ParagraphEmbeddingBlock).filter(ParagraphEmbeddingBlock.document_id == document.id).delete()

        embedding_service = None
        if EmbeddingService.is_configured():
            try:
                embedding_service = EmbeddingService()
            except Exception:
                embedding_service = None

        blocks = self._build_paragraph_blocks(paragraphs)
        created = 0
        for paragraph in paragraphs:
            paragraph.embedding = None
            paragraph.embedding_model = None
            self._db.add(paragraph)

        for block_index, block in enumerate(blocks, start=1):
            embedding_json = None
            embedding_model = None
            if embedding_service is not None:
                started_model = embedding_service.model
                try:
                    result = embedding_service.get_embedding_result(block["text"])
                    embedding_json = json.dumps(result.embedding, separators=(",", ":"))
                    embedding_model = result.model
                    AiLogService(self._db).create(
                        operation="embedding_paragraph_block",
                        status="success",
                        document_id=document.id,
                        model=result.model,
                        request_text=block["text"],
                        response_text=(
                            f"start={block['start_paragraph_index']}; end={block['end_paragraph_index']}; "
                            f"vector_dimensions={len(result.embedding)}"
                        ),
                        prompt_tokens=result.prompt_tokens,
                        total_tokens=result.total_tokens,
                        duration_ms=result.duration_ms,
                    )
                except Exception as exc:
                    embedding_model = started_model
                    AiLogService(self._db).create(
                        operation="embedding_paragraph_block",
                        status="error",
                        document_id=document.id,
                        model=started_model,
                        request_text=block["text"],
                        error_message=str(exc),
                    )

            self._db.add(
                ParagraphEmbeddingBlock(
                    document_id=document.id,
                    block_index=block_index,
                    start_paragraph_index=block["start_paragraph_index"],
                    end_paragraph_index=block["end_paragraph_index"],
                    block_text=block["text"],
                    embedding=embedding_json,
                    embedding_model=embedding_model,
                )
            )
            created += 1

        self._db.commit()
        return created

    def _build_paragraph_blocks(self, paragraphs: list[Paragraph]) -> list[dict[str, int | str]]:
        min_chars = 140
        target_chars = 260
        max_chars = 700
        max_paragraphs = 6

        blocks: list[dict[str, int | str]] = []
        current_texts: list[str] = []
        current_indexes: list[int] = []

        def flush() -> None:
            nonlocal current_texts, current_indexes
            if not current_texts or not current_indexes:
                current_texts = []
                current_indexes = []
                return
            block_text = "\n\n".join(current_texts).strip()
            if block_text:
                blocks.append(
                    {
                        "start_paragraph_index": current_indexes[0],
                        "end_paragraph_index": current_indexes[-1],
                        "text": block_text,
                    }
                )
            current_texts = []
            current_indexes = []

        for paragraph in paragraphs:
            text_value = " ".join((paragraph.text or "").split())
            if not text_value:
                continue

            paragraph_length = len(text_value)
            current_length = sum(len(item) for item in current_texts)
            current_count = len(current_texts)

            if not current_texts:
                current_texts.append(text_value)
                current_indexes.append(int(paragraph.paragraph_index))
                if paragraph_length >= target_chars:
                    flush()
                continue

            should_merge = (
                current_length < min_chars
                or paragraph_length < min_chars
                or current_count < 2
            )
            would_exceed = current_length + paragraph_length > max_chars or current_count >= max_paragraphs

            if not should_merge or would_exceed:
                if current_length >= min_chars:
                    flush()
                elif would_exceed:
                    flush()

            current_texts.append(text_value)
            current_indexes.append(int(paragraph.paragraph_index))

            current_length = sum(len(item) for item in current_texts)
            if current_length >= target_chars and paragraph_length >= min_chars:
                flush()

        flush()
        return blocks

    def load_document_from_path(
        self,
        *,
        file_path: str,
        title: str | None = None,
        author: str | None = None,
        publisher: str | None = None,
        publication_year: int | None = None,
        doc_type: str | None = None,
        bibliography: str | None = None,
    ) -> tuple[Document, dict[str, int]]:
        if self._db is None:
            raise RuntimeError("Database session is required for load_document_from_path().")
        source_path = Path(file_path)
        if not source_path.exists() or not source_path.is_file():
            raise ValueError("Файл ёфт нашуд.")

        file_type, stored_path = self.copy_file_to_storage(source_path)
        document = self.create_uploaded_document(
            filename=source_path.name,
            file_type=file_type,
            stored_path=stored_path,
            title=title,
            author=author,
            publisher=publisher,
            publication_year=publication_year,
            doc_type=doc_type,
            bibliography=bibliography,
        )
        counts = self.parse_document(document)
        self._db.refresh(document)
        return document, counts


def run_ai_processing_job(document_id: int) -> None:
    db = SessionLocal()
    try:
        DocumentParserService(db).run_ai_processing_for_document(document_id=document_id)
    finally:
        db.close()


def start_ai_processing_job(document_id: int) -> None:
    thread = threading.Thread(
        target=run_ai_processing_job,
        args=(document_id,),
        name=f"ai-processing-document-{document_id}",
        daemon=True,
    )
    thread.start()
